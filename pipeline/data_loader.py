"""
data_loader.py

Loads documents from a directory for the Hybrid RAG pipeline.
Supports: PDF, CSV, Excel (.xlsx/.xls), images with text (OCR), TXT, DOCX.

Each loaded file is converted into one or more LangChain `Document` objects
with metadata (source path, file type) attached, ready for chunking/embedding.

Usage:
    from data_loader import load_documents

    docs = load_documents("./data")
"""

import os
import logging
from pathlib import Path
from typing import List

from dotenv import load_dotenv
from langchain_core.documents import Document

# Loaders
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    UnstructuredExcelLoader,
    CSVLoader,
)
from docx import Document as DocxDocument
import pytesseract
from PIL import Image

# Load environment variables (API keys etc.)
load_dotenv()

# ---------------------------------------------------------------------------
# Logging setup
# ---------------------------------------------------------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Supported extensions
# ---------------------------------------------------------------------------
SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".csv",
    ".xlsx",
    ".xls",
    ".txt",
    ".docx",
    ".png",
    ".jpg",
    ".jpeg",
    ".bmp",
    ".tiff",
}

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tiff"}


# ---------------------------------------------------------------------------
# Individual file-type loaders
# ---------------------------------------------------------------------------
def _load_pdf(file_path: Path) -> List[Document]:
    """Load a PDF file page by page."""
    loader = PyPDFLoader(str(file_path))
    docs = loader.load()
    for d in docs:
        d.metadata["source"] = str(file_path)
        d.metadata["file_type"] = "pdf"
    return docs


def _load_txt(file_path: Path) -> List[Document]:
    """Load a plain text file."""
    loader = TextLoader(str(file_path), encoding="utf-8")
    docs = loader.load()
    for d in docs:
        d.metadata["source"] = str(file_path)
        d.metadata["file_type"] = "txt"
    return docs


def _load_csv(file_path: Path) -> List[Document]:
    """Load a CSV file, one Document per row."""
    loader = CSVLoader(str(file_path))
    docs = loader.load()
    for d in docs:
        d.metadata["source"] = str(file_path)
        d.metadata["file_type"] = "csv"
    return docs


def _load_excel(file_path: Path) -> List[Document]:
    """Load an Excel file (.xlsx/.xls)."""
    loader = UnstructuredExcelLoader(str(file_path), mode="elements")
    docs = loader.load()
    for d in docs:
        d.metadata["source"] = str(file_path)
        d.metadata["file_type"] = "excel"
    return docs


def _load_docx(file_path: Path) -> List[Document]:
    """Load a DOCX file by extracting paragraph text."""
    doc = DocxDocument(str(file_path))
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    if not full_text.strip():
        logger.warning(f"No extractable text found in DOCX: {file_path}")
        return []

    return [
        Document(
            page_content=full_text,
            metadata={"source": str(file_path), "file_type": "docx"},
        )
    ]


def _load_image_ocr(file_path: Path) -> List[Document]:
    """Extract text from an image using OCR (pytesseract)."""
    try:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
    except Exception as e:
        logger.error(f"OCR failed for {file_path}: {e}")
        return []

    if not text.strip():
        logger.warning(f"No text extracted via OCR from: {file_path}")
        return []

    return [
        Document(
            page_content=text,
            metadata={"source": str(file_path), "file_type": "image_ocr"},
        )
    ]


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------
_LOADER_MAP = {
    ".pdf": _load_pdf,
    ".txt": _load_txt,
    ".csv": _load_csv,
    ".xlsx": _load_excel,
    ".xls": _load_excel,
    ".docx": _load_docx,
}


def _load_single_file(file_path: Path) -> List[Document]:
    """Route a single file to the correct loader based on its extension."""
    ext = file_path.suffix.lower()

    try:
        if ext in IMAGE_EXTENSIONS:
            return _load_image_ocr(file_path)
        elif ext in _LOADER_MAP:
            return _LOADER_MAP[ext](file_path)
        else:
            logger.warning(f"Unsupported file type skipped: {file_path}")
            return []
    except Exception as e:
        logger.error(f"Failed to load {file_path}: {e}")
        return []


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------
def load_documents(directory: str, recursive: bool = True) -> List[Document]:
    """
    Load all supported documents from a directory.

    Args:
        directory: Path to the folder containing source files.
        recursive: If True, walks into subfolders too.

    Returns:
        List of LangChain Document objects with metadata populated.
    """
    dir_path = Path(directory)
    if not dir_path.exists() or not dir_path.is_dir():
        raise FileNotFoundError(f"Directory not found: {directory}")

    pattern = "**/*" if recursive else "*"
    all_files = [f for f in dir_path.glob(pattern) if f.is_file()]

    relevant_files = [f for f in all_files if f.suffix.lower() in SUPPORTED_EXTENSIONS]

    logger.info(f"Found {len(relevant_files)} supported file(s) in '{directory}'")

    all_documents: List[Document] = []
    for file_path in relevant_files:
        logger.info(f"Loading: {file_path.name}")
        docs = _load_single_file(file_path)
        all_documents.extend(docs)

    logger.info(f"Loaded {len(all_documents)} document chunk(s) total.")
    return all_documents


# ---------------------------------------------------------------------------
# Manual test run
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    DATA_DIR = os.getenv("DATA_DIR", "./data")
    documents = load_documents(DATA_DIR)

    for i, doc in enumerate(documents[:5]):
        print(f"\n--- Document {i+1} ---")
        print(f"Source: {doc.metadata.get('source')}")
        print(f"Type: {doc.metadata.get('file_type')}")
        print(f"Content preview: {doc.page_content[:200]}")